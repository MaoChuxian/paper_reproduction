# 00_extract_and_audit.py
# 提取 Supporting Information 中的 S1-S4，并进行原始数据审计

from pathlib import Path

import pandas as pd

from docx import Document



data_path = Path(
    "data/supporting_information.docx"
)

results_dir = Path(
    "results"
)

results_dir.mkdir(
    exist_ok=True
)



doc = Document(
    data_path
)


print(
    "Word 中表格总数：",
    len(doc.tables)
)



def word_table_to_dataframe(
    table
):

    # 第一行作为表头
    columns = [

        cell.text
        .strip()
        .replace("\n", " ")

        for cell in table.rows[0].cells
    ]


    # 其余行作为数据
    data = []

    for row in table.rows[1:]:

        values = [

            cell.text
            .strip()
            .replace("\n", " ")

            for cell in row.cells
        ]

        data.append(
            values
        )


    return pd.DataFrame(
        data,
        columns=columns
    )



table_s1 = word_table_to_dataframe(
    doc.tables[-4]
)

table_s2 = word_table_to_dataframe(
    doc.tables[-3]
)

table_s3 = word_table_to_dataframe(
    doc.tables[-2]
)

table_s4 = word_table_to_dataframe(
    doc.tables[-1]
)


print()

print(
    "Table S1:",
    table_s1.shape
)

print(
    "Table S2:",
    table_s2.shape
)

print(
    "Table S3:",
    table_s3.shape
)

print(
    "Table S4:",
    table_s4.shape
)



new_columns = [

    "Name",
    "a",
    "b",
    "c",
    "V",
    "concentration",
    "Layer",
    "valence_electron",
    "IQE",
    "Lifetime",
    "Ref"

]


for df in [

    table_s1,
    table_s2,
    table_s3,
    table_s4

]:

    df.columns = new_columns


#
# 注意：
# 此时保留百分数尺度。
#
# 38.9% -> 38.9
#
# 暂时不要变成 0.389。

for df in [

    table_s1,
    table_s2,
    table_s3,
    table_s4

]:

    df["IQE"] = (

        df["IQE"]

        .str.replace(
            "%",
            "",
            regex=False
        )

    )



numeric_columns = [

    "a",
    "b",
    "c",
    "V",
    "concentration",
    "Layer",
    "valence_electron",
    "IQE",
    "Lifetime"

]


for df in [

    table_s1,
    table_s2,
    table_s3,
    table_s4

]:

    for column in numeric_columns:

        df[column] = pd.to_numeric(

            df[column],

            errors="coerce"

        )



table_s1["dopant"] = "Mn4+"

table_s2["dopant"] = "Eu3+"



raw_data = pd.concat(

    [
        table_s1,
        table_s2
    ],

    ignore_index=True

)



print()
print(
    "=" * 60
)

print(
    "原始数据审计"
)

print(
    "=" * 60
)


print()

print(
    "总样本数：",
    len(raw_data)
)


print()

print(
    "Mn4+ 样本数：",
    len(table_s1)
)


print(
    "Eu3+ 样本数：",
    len(table_s2)
)


# 缺失值

print()

print(
    "各列缺失值："
)

print(

    raw_data
    .isnull()
    .sum()

)


# 重复行

print()

print(
    "完全重复样本数：",
    raw_data
    .duplicated()
    .sum()
)


# 描述统计

print()

print(
    "数值特征描述统计："
)

print(

    raw_data[
        numeric_columns
    ]
    .describe()
    .T

)



output_path = results_dir / "00_raw_data.xlsx"


with pd.ExcelWriter(
    output_path
) as writer:

    table_s1.to_excel(
        writer,
        sheet_name="S1_Mn4",
        index=False
    )

    table_s2.to_excel(
        writer,
        sheet_name="S2_Eu3",
        index=False
    )

    table_s3.to_excel(
        writer,
        sheet_name="S3_Test",
        index=False
    )

    table_s4.to_excel(
        writer,
        sheet_name="S4_Tb3",
        index=False
    )

    raw_data.to_excel(
        writer,
        sheet_name="S1_S2_Combined",
        index=False
    )


print()

print(
    "保存完成：",
    output_path
)